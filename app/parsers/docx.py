from __future__ import annotations

import asyncio
from io import BytesIO

from docx import Document

from app.domain.german import normalize_german_text
from app.domain.models import ParsedDocument, ParsedPage, ProcessingStatus


class DocxParser:
    version = "docx_parser:v1"
    supported_mime_types = frozenset(
        {"application/vnd.openxmlformats-officedocument.wordprocessingml.document"}
    )

    async def parse(self, content: bytes, *, file_name: str) -> ParsedDocument:
        text = await asyncio.to_thread(self._extract, content)
        normalized = normalize_german_text(text)
        return ParsedDocument(
            pages=[
                ParsedPage(
                    page_number=1,
                    text=text,
                    char_count=len(text.strip()),
                    has_text_layer=True,
                )
            ],
            original_text=text,
            normalized_text=normalized,
            status=ProcessingStatus.READY if normalized else ProcessingStatus.FAILED,
            parser_version=self.version,
        )

    @staticmethod
    def _extract(content: bytes) -> str:
        document = Document(BytesIO(content))
        paragraphs = [paragraph.text for paragraph in document.paragraphs]
        for table in document.tables:
            for row in table.rows:
                paragraphs.append(" | ".join(cell.text for cell in row.cells))
        return "\n".join(paragraphs)
